#读取word
import docx
doc = docx.Document("demo.docx")
len(doc.paragraphs)
doc.paragraphs[0].text
c
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text