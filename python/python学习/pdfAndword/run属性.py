#run属性
import docx
doc = docx.Document("demo.docx")
doc.paragraphs[0].style
doc.paragraphs[0].style = "Normal"
doc.paragraphs[1].runs[0].style = "QuoteChar"
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save("restyle.docx")