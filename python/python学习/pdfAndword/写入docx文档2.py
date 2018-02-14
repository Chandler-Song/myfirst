#写入docx文档2
import docx
doc = docx.Document()
doc.add_paragraph("Hello world!")
paraObj1 = doc.add_paragraph("This is the second paragraph.")
paraObj2 = doc.add_paragraph("This is the thrid paragraph.")
paraObj1.add_run("This text is being added to the second paragraph.")
doc.save("multi.docx")