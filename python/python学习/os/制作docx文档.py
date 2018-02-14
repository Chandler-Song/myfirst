from docx import Document
from docx.shared import Inches

document = Document()#新建文档

document.add_heading("Documnet Title", 0)

p = document.add_paragraph("A plain paragraph having some")
p.add_run("bold").bold = True#设置字体
p.add_run("and some")
p.add_run('italic.').italic = True

document.add_paragraph("Heading, level 1")#后面加参数有问题style
document.add_paragraph("Intense quote")

document.add_paragraph("first item in unordered list")
document.add_paragraph("first item in unordered list")
document.add_picture("1.png", 1.25, 1.25)

table = document.add_table(rows = 1, cols = 3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
recordset = [[1, 101, "Spam"], [2, 42, "Eggs"], [3, 613, "dogs,pigs"]]
for item in recordset:
	row_cells = table.add_row().cells#增加一行
	row_cells[0].text = str(item[0])#加入到这行中
	row_cells[1].test = str(item[1])
	row_cells[2].text = str(item[2])

document.add_page_break()

document.save("1.docx")